"""Export a TranscriptDocument to various formats.

DOCX and PDF exports optionally accept ``rich_runs`` — per-segment
formatting captured from the editor — so user-applied B/I/U survives the
trip out to attorneys / PI partners. TXT and JSON ignore formatting since
they have no way to represent it.

Language tags and translations
------------------------------
Segments whose ``language`` field is non-empty render with an uppercase
``(XX)`` tag next to the speaker. Translations (hand-written by the
investigator) render as an indented continuation line below the main
line. Rich formatting on the translation is carried through the same
per-segment run lists as the main text, via the ``translation_runs``
parameter on the DOCX and PDF exporters.
"""
from __future__ import annotations

import json
from pathlib import Path
from typing import Optional

from models import FormattedRun, TranscriptDocument, fmt_timestamp


def _speaker_prefix(speaker: str, language: str) -> str:
    """Return the ``"Speaker (ES):"`` / ``"Speaker:"`` / ``""`` prefix
    used in plain-text renderings."""
    tag = f" ({language.upper()})" if language else ""
    if speaker:
        return f"{speaker}{tag}:"
    if language:
        return f"({language.upper()}):"
    return ""


def export_txt(doc: TranscriptDocument, path: Path) -> None:
    lines = []
    if doc.audio_path:
        lines.append(f"# {doc.audio_path.name}")
    if doc.language:
        lines.append(f"# Language: {doc.language} ({doc.language_probability:.2f})")
    lines.append("")
    for seg in doc.segments:
        prefix = _speaker_prefix(seg.speaker, seg.language)
        head = f"[{fmt_timestamp(seg.start)} -> {fmt_timestamp(seg.end)}]"
        if prefix:
            lines.append(f"{head} {prefix} {seg.text}")
        else:
            lines.append(f"{head} {seg.text}")
        if seg.translation:
            # Indented translation line — mirrors the editor's visual layout.
            lines.append(f"    \u21b3 {seg.translation}")
    path.write_text("\n".join(lines), encoding="utf-8")


def export_json(doc: TranscriptDocument, path: Path) -> None:
    path.write_text(
        json.dumps(doc.to_dict(), indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


def _runs_for_segment(
    seg_index: int,
    seg_text: str,
    rich_runs: Optional[list[list[FormattedRun]]],
) -> list[FormattedRun]:
    """Return the formatted runs for a segment, falling back to a single
    plain run when no formatting was captured (e.g., headless export)."""
    if rich_runs and seg_index < len(rich_runs) and rich_runs[seg_index]:
        return rich_runs[seg_index]
    return [FormattedRun(text=seg_text)]


def export_docx(
    doc: TranscriptDocument,
    path: Path,
    rich_runs: Optional[list[list[FormattedRun]]] = None,
    translation_runs: Optional[list[list[FormattedRun]]] = None,
) -> None:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor

    d = DocxDocument()
    if doc.audio_path:
        d.add_heading(doc.audio_path.name, level=1)

    for i, seg in enumerate(doc.segments):
        p = d.add_paragraph()
        ts_run = p.add_run(f"[{fmt_timestamp(seg.start)} -> {fmt_timestamp(seg.end)}]  ")
        ts_run.bold = True
        ts_run.font.size = Pt(9)
        ts_run.font.color.rgb = RGBColor(108, 117, 125)
        # Speaker prefix — appends the "(ES)" language tag when the
        # segment isn't in the document's default language.
        tag = f" ({seg.language.upper()})" if seg.language else ""
        if seg.speaker:
            spk_run = p.add_run(f"{seg.speaker}{tag}: ")
            spk_run.bold = True
            spk_run.font.size = Pt(11)
            spk_run.font.color.rgb = RGBColor(0, 102, 204)
        elif seg.language:
            # Rare: language tag but no speaker yet — still show the tag.
            spk_run = p.add_run(f"({seg.language.upper()}): ")
            spk_run.bold = True
            spk_run.font.size = Pt(11)
            spk_run.font.color.rgb = RGBColor(0, 102, 204)

        for run in _runs_for_segment(i, seg.text, rich_runs):
            text_run = p.add_run(run.text)
            text_run.font.size = Pt(11)
            if run.bold:
                text_run.bold = True
            if run.italic:
                text_run.italic = True
            if run.underline:
                text_run.underline = True

        # Translation continuation paragraph — indented, italicised by
        # default so it's visually distinct from the spoken words. The
        # investigator's own B/I/U overrides layer on top.
        if seg.translation:
            tp = d.add_paragraph()
            tp.paragraph_format.left_indent = Pt(24)
            arrow = tp.add_run("\u21b3  ")
            arrow.font.size = Pt(11)
            arrow.font.color.rgb = RGBColor(120, 120, 140)
            trans_runs = _runs_for_segment(i, seg.translation, translation_runs)
            for run in trans_runs:
                tr = tp.add_run(run.text)
                tr.font.size = Pt(11)
                # Italic by default for translations. The user can
                # override by explicitly applying (or removing) B/I/U
                # in the editor — those flags come through in the run.
                tr.italic = True
                tr.bold = run.bold
                if run.underline:
                    tr.underline = True
                tr.font.color.rgb = RGBColor(90, 90, 110)

    d.save(str(path))


def export_pdf(
    doc: TranscriptDocument,
    path: Path,
    rich_runs: Optional[list[list[FormattedRun]]] = None,
    translation_runs: Optional[list[list[FormattedRun]]] = None,
) -> None:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Title
    if doc.audio_path:
        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(0, 12, doc.audio_path.name, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    if doc.language:
        pdf.set_font("Helvetica", "I", 9)
        pdf.cell(0, 6, f"Language: {doc.language} ({doc.language_probability:.2f})", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    for i, seg in enumerate(doc.segments):
        ts = f"[{fmt_timestamp(seg.start)} -> {fmt_timestamp(seg.end)}]"

        # Timestamp
        pdf.set_font("Helvetica", "B", 8)
        pdf.set_text_color(108, 117, 125)
        pdf.cell(pdf.get_string_width(ts) + 2, 6, ts)

        # Speaker label (appends "(ES)" tag when segment isn't default lang)
        tag = f" ({seg.language.upper()})" if seg.language else ""
        spk_label = ""
        if seg.speaker:
            spk_label = f"  {seg.speaker}{tag}: "
        elif seg.language:
            spk_label = f"  ({seg.language.upper()}): "
        if spk_label:
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(0, 102, 204)
            pdf.cell(pdf.get_string_width(spk_label) + 2, 6, spk_label)

        # Text — emit each formatted run as a separate write() so B/I/U
        # styling sticks. Underline isn't a font-style flag in fpdf2; it's
        # part of the style string ("U", "BU", "BIU", etc.).
        pdf.set_text_color(0, 0, 0)
        runs = _runs_for_segment(i, seg.text, rich_runs)
        prefix = "" if spk_label else "  "
        for j, run in enumerate(runs):
            style = ""
            if run.bold:
                style += "B"
            if run.italic:
                style += "I"
            if run.underline:
                style += "U"
            pdf.set_font("Helvetica", style, 10)
            text = (prefix + run.text) if j == 0 else run.text
            pdf.write(6, text)
        pdf.ln(8)

        # Translation continuation (indented, italicised by default)
        if seg.translation:
            pdf.set_text_color(90, 90, 110)
            pdf.set_font("Helvetica", "I", 10)
            pdf.write(6, "        \u21b3  ")
            trans_runs = _runs_for_segment(i, seg.translation, translation_runs)
            for run in trans_runs:
                style = "I"  # italic by default for translation body
                if run.bold:
                    style += "B"
                if run.underline:
                    style += "U"
                pdf.set_font("Helvetica", style, 10)
                pdf.write(6, run.text)
            pdf.ln(8)

        pdf.ln(1)

    pdf.output(str(path))
